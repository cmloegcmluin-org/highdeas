"""Routers that hand a submitted memo to its chosen destination.

Three of them deliver it outright — Notesnook, Google Drive, Asana. The fourth
opens it in Claude as a prompt nobody has sent yet."""
import html
import re
import shutil
import urllib.error
from datetime import datetime
from pathlib import Path
from urllib.parse import quote, urlencode

import requests
from asana_client import create_subtask

# The module, not the class: drive_write is re-executed on purpose (see
# test_the_engine_starts_on_a_machine_without_google_auth, which imports it as a
# machine missing google-auth would), and a reload mints a new exception class.
# A name bound here at import time would then be a stale class this `except`
# quietly stops matching. Looked up at raise time, both sides always agree.
from highdeas import drive_write
from highdeas.nonspeech import UNCLEAR


# A note is stored as plain text, so a list is just its Markdown line — which is
# what the editor writes and reads back. Both destinations turn those lines into
# real lists: HTML for Notesnook, Word's list styles for a Drive .docx.
_BULLET = re.compile(r"^\s*[-*•]\s+(.*)$")
_NUMBER = re.compile(r"^\s*\d+[.)]\s+(.*)$")


def _list_item(line):
    """The item text and its list tag ("ul"/"ol"), or (None, None) for prose."""
    bullet = _BULLET.match(line)
    if bullet:
        return bullet.group(1).strip(), "ul"
    number = _NUMBER.match(line)
    if number:
        return number.group(1).strip(), "ol"
    return None, None


def _text_to_html(text):
    parts = []
    open_tag = None
    for line in text.split("\n"):
        item, tag = _list_item(line)
        if tag != open_tag:
            if open_tag:
                parts.append(f"</{open_tag}>")
            if tag:
                parts.append(f"<{tag}>")
            open_tag = tag
        if tag:
            parts.append(f"<li>{html.escape(item)}</li>")
        elif line.strip():
            parts.append(f"<p>{html.escape(line.strip())}</p>")
    if open_tag:
        parts.append(f"</{open_tag}>")
    return "".join(parts) or "<p></p>"


def _default_title(memo):
    """Name an unnamed memo the way Notesnook names untitled notes ("Note $date$
    $time$"), from when it happened: the moment it was recorded, or failing that
    the moment the app first saw it. Notesnook's Inbox API requires a non-empty
    title, so this always returns one even when neither time is known.

    The time is to the second, not the minute: two unnamed memos recorded in the
    same minute would otherwise share a title, and same-titled notes collapse to one
    in the inbox — silently dropping every second recording made within a minute."""
    try:
        made = datetime.fromisoformat(memo.recorded_at or memo.created_at)
    except (TypeError, ValueError):
        return "Voice note"
    hour = made.hour % 12 or 12
    meridiem = "AM" if made.hour < 12 else "PM"
    return f"Note {made:%Y-%m-%d} {hour}:{made:%M}:{made:%S} {meridiem}"


class NotesnookRouter:
    """Create a note via the Notesnook Inbox API (POST https://inbox.notesnook.com/)."""

    ENDPOINT = "https://inbox.notesnook.com/"

    def __init__(self, api_key, *, source="highdeas", post=requests.post):
        self._api_key = api_key
        self._source = source
        self._post = post

    def route(self, memo):
        response = self._post(
            self.ENDPOINT,
            headers={"Authorization": self._api_key, "Content-Type": "application/json"},
            json={
                "title": memo.name or _default_title(memo),
                "type": "note",
                "source": self._source,
                "version": 1,
                "content": {"type": "html", "data": _text_to_html(memo.transcript)},
            },
            timeout=30,
        )
        response.raise_for_status()


def parse_choices(raw):
    """Read how .env spells a dropdown — ";"-separated "value=Label" pairs — into an
    ordered list of (value, label). The first pair leads the dropdown and is what an
    unanswered note falls back to; a pair without "=Label" is labelled by its value.

    Two dropdowns are configured this way. ASANA_PARENT_TASKS names the tasks a note
    can be filed under, where a gid may carry an "ACCOUNT:" prefix saying which Asana
    account holds it (see _account_and_gid) — the prefix travels as part of the value,
    so the row neither shows it nor has to understand it. HIGHDEAS_CLAUDE_MODELS names
    the models a chat can be opened on, as the ids claude.ai takes in a link."""
    choices = []
    for pair in (raw or "").split(";"):
        value, _, label = pair.partition("=")
        value, label = value.strip(), label.strip()
        if value:
            choices.append((value, label or value))
    return choices


def _account_and_gid(parent):
    """Split a dropdown value into the account holding the task and the task's gid.
    "WORK:333" is task 333 in the "WORK" account; a bare "333" is the account the
    app has always had, whose token is the unsuffixed ASANA_ACCESS_TOKEN."""
    account, marked, gid = parent.partition(":")
    return (account, gid) if marked else ("", parent)


def _asana_token_variable(account):
    """The .env variable holding an account's token — ASANA_ACCESS_TOKEN for the
    unnamed default, ASANA_ACCESS_TOKEN_WORK for "work".

    Upper-cased because that is what the marker names: a variable. Environment
    lookups are case-sensitive on the Mac and not on Windows, so a marker taken
    literally would submit fine at one desk and 401 at the other."""
    return f"ASANA_ACCESS_TOKEN_{account.upper()}" if account else "ASANA_ACCESS_TOKEN"


def read_asana_tokens(parents, env):
    """The access token for every account the offered parents name, read from `env`.
    Each account is one more personal access token; only the accounts actually on
    the dropdown are looked for, so a second one costs nothing until a task names
    it. The default account is always among them, so a submit with nothing
    configured can still name the variable to fill in."""
    accounts = {""} | {_account_and_gid(parent)[0] for parent, _ in parents}
    return {account: env.get(_asana_token_variable(account), "") for account in accounts}


def _group_items(transcript):
    """A group's bullets as (task name, task notes) pairs.

    Groups write "- Name: text" for a named note and "- text" for a bare one (see the
    service's `_bullet`), so a colon splits name from notes; a line without one is all name.
    Numbered lines and stray prose ride in when an absorbed group brought its own transcript
    across, so every line with words on it is an item; blank lines are not."""
    for line in transcript.splitlines():
        item, _ = _list_item(line)
        if item is None:
            item = line.strip()
        if not item:
            continue
        name, _, notes = item.partition(": ")
        yield (name.strip(), notes.strip()) if notes else (item, "")


class AsanaRouter:
    """Create the memo's text as a subtask of its chosen Asana parent task
    (POST /tasks/{gid}/subtasks) — and a GROUP's items as one subtask each. Only
    the text goes to Asana — the note's name and transcript; the recording itself
    stays local and retires to the bin like every other route. Reports the created
    task's permalink for the memo's record, so the bin icon can open the task.

    Holds one token per Asana account, since a parent task names the account it
    belongs to: two accounts are two tokens behind one dropdown. The request
    itself is the family's shared client (asana_client, the asana-mcp repo) -
    Excephalon files tasks through the same code, so the two apps cannot grow
    separate spellings of one request."""

    def __init__(self, tokens, *, default_parent="", create=create_subtask):
        self._tokens = tokens
        self._default_parent = default_parent
        self._create_subtask = create

    def route(self, memo):
        parent = memo.asana_parent or self._default_parent
        if not parent:
            raise RuntimeError("No Asana parent task configured — put ASANA_PARENT_TASKS in .env.")
        account, gid = _account_and_gid(parent)
        token = self._tokens.get(account, "")
        if not token:
            raise RuntimeError("Asana access token not set — put "
                               f"{_asana_token_variable(account)} in .env.")
        # A NAMELESS group fans out: it has no name to preserve, so its bulleted consolidation
        # would otherwise become one task named after the literal bullets — each item becomes
        # its own subtask instead. A NAMED group must NOT fan out: splitting throws the name
        # away, so it falls through to the single-task path below, where the group's name
        # becomes the task and its bullets (or numbered items) ride along as the task's notes.
        if memo.kind == "group" and not memo.name:
            first = ""
            for name, notes in _group_items(memo.transcript):
                created = self._create(gid, token, name, notes)
                first = first or created
            return {"asana_url": first}
        # A named memo — note or group — keeps its transcript as the task's notes. An unnamed
        # one has only its transcript, so that becomes the name — a readable task rather than
        # a generic date title — and the notes are left empty rather than repeating it.
        if memo.name:
            name, notes = memo.name, memo.transcript
        else:
            name = memo.transcript or _default_title(memo)
            notes = ""
        return {"asana_url": self._create(gid, token, name, notes)}

    def _create(self, gid, token, name, notes):
        """One subtask via the shared client; the permalink comes back for the memo's record.

        A refusal carries Asana's own words ("Not a recognized ID", "You do not have
        access") - an opaque status code sends whoever reads the error off to fix the
        wrong thing."""
        try:
            return self._create_subtask(token, gid, name, notes)
        except urllib.error.HTTPError as denied:
            words = denied.read().decode("utf-8", errors="replace")[:300]
            raise RuntimeError(f"Asana refused the task: {words}") from denied


def _link(base, **params):
    """`base` with its non-empty `params` as a query string, spaces as %20 rather
    than "+". An empty value is left out entirely: "model=" with nothing behind it
    is a request for a model named "", not the absence of a request."""
    given = {name: value for name, value in params.items() if value}
    return f"{base}?{urlencode(given, quote_via=quote)}"


class ClaudeRouter:
    """Open the note as an unsent prompt in a new Claude session.

    Nothing is sent: both surfaces fill the composer and stop, so the memo leaves
    Highdeas as a question waiting to be read rather than as a delivered note."""

    def __init__(self, open_browser, open_deep_link, *, folder=""):
        self._open_browser = open_browser
        self._open_deep_link = open_deep_link
        self._folder = folder

    def route(self, memo):
        prompt = "\n\n".join(part for part in (memo.name, memo.transcript) if part)
        if memo.claude_surface == "chat":
            self._open_browser(_link("https://claude.ai/new", q=prompt,
                                     model=memo.claude_model))
        else:
            self._open_deep_link(_link("claude://code/new", q=prompt,
                                       folder=self._folder))


class Router:
    """Dispatch a memo to the router for its chosen route (Notesnook by default),
    passing through whatever fields that router reports for the store to persist
    (e.g. Asana's link to the created task).

    One key in that dict is not a memo field: "warning" is a sentence about a send
    that landed in a lesser form than it should have (see DriveMusicRouter and the
    .docx fallback). InboxService.submit takes it out before the rest is stored and
    hands it back for the page to show, because a degraded send is neither a failure
    to raise nor a success to pass over in silence."""

    def __init__(self, notesnook, drive=None, asana=None, claude=None):
        self._routers = {"notesnook": notesnook, "drive": drive, "asana": asana,
                         "claude": claude}

    def __call__(self, memo):
        router = self._routers.get(memo.route, self._routers["notesnook"])
        if router is not None:
            return router.route(memo)
        return None


DATE_FORMAT = "%Y_%m_%d"


def _today():
    return datetime.now().strftime(DATE_FORMAT)


def drive_subfolder_name(date_str):
    """The dated Drive subfolder name a music memo is filed into, given a
    DATE_FORMAT date string. Pulled out of DriveMusicRouter.route() as its own
    function so anything reconstructing a *past* subfolder name — the backfill
    script (scripts/backfill_drive_subfolders.py) for a memo routed before
    drive_subfolder was tracked — can call the exact same code route() does,
    from a date it recovers from that memo's own stored processed_at, instead
    of maintaining a second copy of this format that could drift out of sync."""
    return f"_{date_str}_NOT_YET_PROCESSED_MUSIC"


def _sanitize_filename(name):
    cleaned = re.sub(r'[<>:"/\\|?*]', "", name).strip()
    return cleaned or "untitled"


_LIST_STYLE = {"ul": "List Bullet", "ol": "List Number"}


def write_docx(path, text):
    from docx import Document

    document = Document()
    for line in text.split("\n"):
        item, tag = _list_item(line)
        if tag:
            document.add_paragraph(item, style=_LIST_STYLE[tag])
        else:
            document.add_paragraph(line)
    document.save(str(path))


def _says_nothing(transcript):
    """Whether a transcript has no words of its own to file: blank, or nothing but the
    "[unclear]" the transcriber writes where it heard no speech at all (nonspeech.py) —
    a single line of it, or, for a group of takes it could read none of, one per bullet
    (service._bullet). A doc of those says nothing the memo itself does not already say.

    A bullet that names its item ("- Chorus: [unclear]") does say something, and is not
    this — the name is content the doc is the only place to keep."""
    for line in transcript.splitlines():
        item, _tag = _list_item(line)
        text = (item if item is not None else line).strip()
        if text and text != UNCLEAR:
            return False
    return True


class DriveMusicRouter:
    """Copy a music memo into a dated folder under the Drive base, with an optional doc.

    The original stays in the inbox so the service can also retire it to the local
    bin — the memo is then recoverable there for 90 days regardless of what happens
    to the Drive copy.

    The doc itself is filed two different ways depending on what's configured. When
    `file_doc` is given (see app._drive_doc_filer) it's tried first: a real, native
    Google Doc created through the actual Drive API — no more Word-into-a-synced-
    folder trick. Only when that isn't configured, or the call comes back empty (not
    authorized yet, offline, any Drive hiccup), does the old local .docx write run —
    the transcript must reach *some* filed doc rather than none. A native Doc always
    files into its own container folder first (see drive_write.py's module docstring
    for why), then moves beside the audio itself when `file_doc` can resolve it —
    immediately, most of the time, or later via DriveDocReconciler when it can't yet.
    `drive_doc_link` on the returned outcome is the way back to the doc either way;
    `drive_doc_needs_move` says whether that move is still pending.

    Neither runs for a transcript with nothing of its own to say (see `_says_nothing`):
    the audio is filed alone rather than beside a doc reading only "[unclear]"."""

    def __init__(self, inbox_dir, drive_base, *, today=_today, write_doc=write_docx,
                 file_doc=None, copy=shutil.copy2, wake_drive=None):
        self._inbox = Path(inbox_dir)
        self._base = Path(drive_base)
        self._today = today
        self._write_doc = write_doc
        self._file_doc = file_doc
        self._copy = copy
        # Called with the base folder when it's missing, to start Drive for Desktop
        # and wait briefly for the mount (app.py wires drive_mount.wake). Injected
        # rather than defaulted, so that a router built in a test never launches a
        # desktop app on the machine running the suite.
        self._wake_drive = wake_drive

    def route(self, memo):
        # The dated subfolder is ours to create; the base underneath it is Drive's,
        # and its absence means Drive is not here — not installed, not signed in, or
        # HIGHDEAS_DRIVE_BASE still carrying the other machine's path. Making it
        # would put the memo in a folder that merely looks like Drive, on a disk
        # nothing uploads from, and report the send as done. Refuse instead: the
        # submit route turns this into a sentence in the notice bar and the memo
        # stays in the inbox, where it can be sent again once Drive is really there.
        #
        # One thing is worth trying before that refusal, though, because it is the
        # commonest cause by far on a machine that does have Drive: Drive for Desktop
        # not running. wake_drive starts it and waits a bounded few seconds — and
        # then the folder is read again, so it, not the attempt, is still what says
        # whether this memo may be filed.
        if not self._base.is_dir() and self._wake_drive is not None:
            self._wake_drive(self._base)
        if not self._base.is_dir():
            raise FileNotFoundError(
                f"Google Drive folder not found: {self._base} — check that Drive for "
                f"Desktop is running and that HIGHDEAS_DRIVE_BASE in .env names this "
                f"machine's own Drive path."
            )
        subfolder_name = drive_subfolder_name(self._today())
        folder = self._base / subfolder_name
        folder.mkdir(parents=True, exist_ok=True)
        source = self._inbox / memo.audio_filename
        base = _sanitize_filename(memo.name or Path(memo.audio_filename).stem)
        self._copy(str(source), str(folder / (base + source.suffix)))
        doc_link, needs_move, warning = "", False, ""
        if not _says_nothing(memo.transcript):
            if self._file_doc is not None:
                title = memo.name or _default_title(memo)
                try:
                    doc_link, needs_move = self._file_doc(
                        subfolder_name, title, _text_to_html(memo.transcript))
                except drive_write.DriveDocUnavailable as exc:
                    # The docx below still runs -- the transcript reaches a filed doc
                    # either way -- but a Docs setup that has stopped working is now
                    # said out loud on the first submit that hits it, rather than
                    # discovered later by noticing the file format.
                    warning = (f"Filed the transcript as a .docx instead of a Google "
                               f"Doc: {exc}")
            if not doc_link:
                self._write_doc(folder / (base + ".docx"), memo.transcript)
        # Nothing here yet knows this subfolder's own Drive ID — Drive for Desktop
        # uploads it to the cloud on its own schedule — so only its name is reported;
        # the bin's Drive icon looks up the ID from this name later, via the real
        # Drive API, when it's asked to open this memo.
        return {"drive_subfolder": subfolder_name, "drive_doc_link": doc_link,
                "drive_doc_needs_move": needs_move, "warning": warning}
